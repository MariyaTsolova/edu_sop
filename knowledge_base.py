# %%
import os
import json
import pickle
import streamlit as st
from random import randrange
import torch

from haystack import Pipeline
from haystack import Document
from haystack.core.component import Component
from haystack import component
from haystack.document_stores.in_memory import InMemoryDocumentStore
from haystack_integrations.document_stores.chroma import ChromaDocumentStore

# from haystack.components.converters import PDFMinerToDocument
from haystack.components.converters import PyPDFToDocument, TextFileToDocument  #  After minimal Testing, PyPDFToDocument appears to be working better for our documents 
from haystack.components.embedders import SentenceTransformersDocumentEmbedder

from haystack.components.preprocessors import DocumentCleaner
from haystack.components.preprocessors import DocumentSplitter
from haystack.components.writers import DocumentWriter
# from haystack.components.preprocessors import RecursiveDocumentSplitter

import contextualiser

import docx
import chromadb
from chromadb.utils import embedding_functions
import uuid

from sentence_transformers import SentenceTransformer
from haystack.components.retrievers.in_memory import InMemoryBM25Retriever
from haystack.document_stores.in_memory import InMemoryDocumentStore


torch.classes.__path__ = []

document_store = ChromaDocumentStore("SOP-Docs-LocalEMB", host="localhost", port=8077)
client = chromadb.HttpClient(port='8077')
model_name = 'all-mpnet-base-v2'

#ChormaDB client and doc_store
doc_store = client.get_collection("SOP-Docs-LocalEMB_v2")
doc_store_cs = client.get_collection("SOP_case_studies_embed_full_v4")

model = SentenceTransformer('sentence-transformers/all-mpnet-base-v2')

def query_knowledge_base(query_text, n=5):
    embedding_model = SentenceTransformer('sentence-transformers/all-mpnet-base-v2')
    embedding = embedding_model.encode(query_text)
    res = doc_store.query(embedding, n_results=n)
    return res

def query_knowledge_base_cs(query_text, n=5):
    embedding_model = SentenceTransformer('sentence-transformers/all-mpnet-base-v2')
    embedding = embedding_model.encode(query_text)
    res = doc_store_cs.query(embedding, n_results=n)
    return res

def format_chunks(documents):
    chunks = documents['documents'][0]
    metas = documents['metadatas'][0]
    distances = documents['distances'][0]

    # chunks_all_info = [{"content": c,
    #                     "filepath": m['file_path'],
    #                     "page_number": m['page_number'],
    #                     "URL": m['url'],
    #                     "cos_dist": d} for c, m, d in zip(chunks, metas, distances)]
    
    chunks_all_info = []
    for c, m, d in zip(chunks, metas, distances):

        chunk = f"""
        Source: {m['file_path']}
        URL: {m['url']}
        Page Number: {m['page_number']}
        Cosine Distance: {d}
        ---
        {c}
        """
        chunks_all_info.append(chunk.strip())
    
    return "\n\n".join(chunks_all_info)


def format_chunks_cs(documents):
    content_cs = [res.content for res in documents]
    meta_cs = [res.meta for res in documents]
    # ids_cs = [res.id for res in res_cs]
    score_cs = [res.score for res in documents]

    # chunks_all_info = [{"content": c,
    #                     "filepath": m['file'].split('/')[-1],
    #                     "page_number": m['table_in_file'],
    #                     "URL": 'Case study by Educational Psychologist',
    #                     "cos_dist": d} for c, m, d in zip(content_cs, meta_cs, score_cs)]
    
    chunks = []
    for c, m, d in zip(content_cs, meta_cs, score_cs):

        chunk = f"""
        File Name: {m['file'].split('/')[-1]}
        Source: Case study by Educational Psychologist
        Table in file: {m['table_in_file']}
        Cosine Distance: {d}
        ---
        {c}
        """
        chunks.append(chunk.strip())
    
    return "\n\n".join(chunks)


def init_bm25_retriever(chroma_collection):
    chroma_data = chroma_collection.get()
    docs = [
        Document(content=doc, meta=meta, id=id_)
        for doc, meta, id_ in zip(chroma_data['documents'], chroma_data['metadatas'], chroma_data['ids'])
    ]

    doc_store = InMemoryDocumentStore()
    doc_store.write_documents(documents=docs)
    bm25_retriever = InMemoryBM25Retriever(document_store=doc_store)
    return bm25_retriever

bm25_retriever_cs = init_bm25_retriever(doc_store_cs)

def hybrid_search_cs(query, top_k = 5, alpha = 0.5):
    bm25_results = bm25_retriever_cs.run(query=query, top_k=top_k, scale_score=True)["documents"]

    dense_results_raw = query_knowledge_base_cs(query)

    dense_results = [
        Document(content=text, meta=meta, id=id_, score=score)
        for text, meta, id_, score in zip(
            dense_results_raw['documents'][0],
            dense_results_raw['metadatas'][0],
            dense_results_raw['ids'][0],
            dense_results_raw['distances'][0]
        )
    ]

    combined = {}
    for doc in bm25_results:
        combined[doc.id] = {'doc':doc, 'score':(1-alpha)*doc.score}
    for doc in dense_results:
        combined.setdefault(doc.id, {'doc':doc, 'score':0})
        combined[doc.id]['score'] += alpha * doc.score

    sorted_docs = sorted(combined.values(), key = lambda x: x['score'], reverse = True)

    return [entry['doc'] for entry in sorted_docs[:top_k]]


     





# ===============================================================================
# Until here Functions for using Knowledgebase <<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<
#
# >>>>>>>>>>>>>>>>>>>>>>>>> From here fucntion for creating Knowledgebase
# ================================================================================

CHUNK_SIZE_WORD = 250
CHUNK_SIZE_SENT = 5

doc_class = "compliance" 
# path_choma_document_store = os.path.join("data", "chroma", "documents_KB")
data_folder = os.path.abspath(os.path.join("data", "knowledge_base_documents_sop"))

if os.path.isdir(data_folder):
    docs_folder = os.path.join(data_folder)

    docs_names = os.listdir(docs_folder)
    docs_paths = [os.path.join(docs_folder, d) for d in docs_names]

    meta_docs_urls = [{'url': 'doc_not_public_yet',  "document class": doc_class}] * len(docs_names)

    pdf_kb = False
    sc_kb = False
    comb_kb = False
    case_studies_kb = True
# %%
""" Testing
converter = PyPDFToDocument()
docs = converter.run(sources=docs_paths[4:6])['documents']

cleaner = DocumentCleaner()
docs_c = cleaner.run(docs)['documents']

# splitter2 = RecursiveDocumentSplitter(split_length=CHUNK_SIZE)
# splitter2.warm_up()
# docs_s = splitter2.run(docs_c)['documents']

splitter = DocumentSplitter(split_by = "word", split_length = CHUNK_SIZE, split_overlap = 25)
docs_s = splitter.run(docs_c)

"""

#%%    
def preprocess_pdfs(document_store):
    print("pdf preprocessing beginning")

    openai_key = st.secrets["API_keys"]["openai"]
    os.environ['OPENAI_API_KEY'] = openai_key

    document_splitter = DocumentSplitter(split_by="sentence", split_length=CHUNK_SIZE_SENT, split_overlap=1)
    document_contextualiser = contextualiser.ContextualTextPreProcessor(chunk_size=CHUNK_SIZE_SENT, model="gpt-4.1-nano", api_key=openai_key)
    document_embedder = SentenceTransformersDocumentEmbedder(f'sentence-transformers/{model_name}')


   
    # TODO - Add bm25 embedding pdf files
    # TODO - It cannot read and convert to Documents, two of the pdfs, cryptography>=3.1 is required for AES algorithm
    document_writer = DocumentWriter(document_store)
    
    pdf_preprocessing_pipeline = Pipeline()
    pdf_preprocessing_pipeline.add_component(instance=PyPDFToDocument(), name="pdf_converter")
    pdf_preprocessing_pipeline.add_component(instance=DocumentCleaner(), name="document_cleaner")
    pdf_preprocessing_pipeline.add_component(instance=document_splitter, name="document_splitter")
    pdf_preprocessing_pipeline.add_component(instance=document_contextualiser, name="document_contextualiser")
    pdf_preprocessing_pipeline.add_component(instance=document_embedder, name="document_embedder")
    pdf_preprocessing_pipeline.add_component(instance=document_writer, name="document_writer")

    pdf_preprocessing_pipeline.connect("pdf_converter", "document_cleaner")
    pdf_preprocessing_pipeline.connect("document_cleaner", "document_splitter")
    pdf_preprocessing_pipeline.connect("document_cleaner", "document_contextualiser.og_docs")
    pdf_preprocessing_pipeline.connect("document_splitter", "document_contextualiser.documents")
    pdf_preprocessing_pipeline.connect("document_contextualiser", "document_embedder")
    pdf_preprocessing_pipeline.connect("document_embedder", "document_writer")
    pdf_preprocessing_pipeline.run({"pdf_converter": {"sources": docs_paths, "meta": meta_docs_urls}})
    print("pdf preprocessing pipeline finished")

    return document_store

def change_embedding_collection(client, old_collection_name, new_collection_name, model_name):

    embedding_func = embedding_functions.SentenceTransformerEmbeddingFunction(model_name = f"sentence-transformers/{model_name}")

    oldColl = client.get_collection(old_collection_name)

    newColl = client.create_collection(name=new_collection_name, 
                                       embedding_function=embedding_func, 
                                       metadata={'embedding_model':model_name})

    existing_count = oldColl.count()
    batch_size = 1000
    for i in range(0, existing_count, batch_size):
        batch = oldColl.get(include=["metadatas", "documents"], limit=batch_size, offset=i)
        newColl.add(
            ids=batch["ids"],
            documents=batch["documents"],
            metadatas=batch["metadatas"])

    return newColl

def input_scenario_to_string_variation_2(s):
    string = f"Student Profile:\n{json.dumps(s['student_profile'])}\n\nSituation:\n{s['situation']}\n\nAction Taken:\n{s['action']}"
    return string

def convert_all_json_to_text():
    path_scenario_folder = os.path.join("data", "new_synthetic_scenarios_v2")
    paths_scenarios = [os.path.join(path_scenario_folder, s) for s in os.listdir(path_scenario_folder)]
    scenarios = []
    for path in paths_scenarios:
        with open(path, "r") as f:
            s = json.load(f)
            s['file_path'] = path
        
        meta = {
                    "file": s['file_path'],
                    "grade": s['grade'],
                    "effect": s['effect'],
                    "topic": s['topic'],
                    "len": s['len']
                }   
        text = input_scenario_to_string_variation_2(s)

        scenarios.append(Document(content = text, meta=meta))
    return scenarios


def preprocess_scenarios(document_store):

    # TODO Decide what to do with the metadata 
    print("scenarios preprocessing beginning")
    scenarios_txt = convert_all_json_to_text()

    scenario_preprocessing_pipeline = Pipeline()
    json_embedder = SentenceTransformersDocumentEmbedder()
    # TODO - Add bm25 embedding txt files
    document_writer_json = DocumentWriter(document_store)

    scenario_preprocessing_pipeline.add_component(instance=json_embedder, name="json_embedder")
    scenario_preprocessing_pipeline.add_component(instance=document_writer_json, name="document_writer_json")

    scenario_preprocessing_pipeline.connect("json_embedder", "document_writer_json")

    scenario_preprocessing_pipeline.run({"json_embedder": {"documents": scenarios_txt}})
    print("json preprocessing pipeline finished")
    
    return document_store

def get_rand_scenario_high_grade():
    path_scenario_folder = os.path.join("data", "new_synthetic_scenarios_v2")
    paths_scenarios = [os.path.join(path_scenario_folder, s) for s in os.listdir(path_scenario_folder)]
    scenarios_grade_five = []
    for path in paths_scenarios:
        with open(path, "r") as f:
            s = json.load(f)
        if s['grade']==5:
            scenarios_grade_five.append(path)
    scenario_path = scenarios_grade_five[randrange(0, len(scenarios_grade_five))]
    return scenario_path

# %% Combining pdf documents with txt scenarios 

def create_pdf_kb():
    document_store_pdf = preprocess_pdfs(InMemoryDocumentStore(embedding_similarity_function="cosine"))
    return document_store_pdf

def create_scenario_kb():
    document_store_scenario = preprocess_scenarios(InMemoryDocumentStore(embedding_similarity_function="cosine"))
    return document_store_scenario

def create_combined_kb():

    kb_pdfs_pkl_file_path = os.path.join(data_folder, "doc_store_pdfs.pkl")
    if os.path.exists(kb_pdfs_pkl_file_path):
        doc_store_pdfs = InMemoryDocumentStore.load_from_disk(kb_pdfs_pkl_file_path)
    else:
        doc_store_pdfs = preprocess_pdfs(InMemoryDocumentStore(embedding_similarity_function="cosine"))

    document_store_combined = preprocess_scenarios(doc_store_pdfs)
    return document_store_combined

# %%

def parse_docx_to_dict(path):
    document_docx = docx.Document(path)

    # Parse every table into a json/dict of 
    case_studies = []
    for i, table in enumerate(document_docx.tables):
        case_study = {"id": str(uuid.uuid4()), 
                      "case_id_in_doc": str(i),
                      "docx_path": path}

        # We semantically divide the table into header and plan/content
        in_header = True
        rows = []
        for row in table.rows:
            contents = [cell.text.strip() for cell in row.cells]
            unique = []
            [unique.append(x) for x in contents if x not in unique]
            l = len(unique)
            if  l < 2:
                print("Unexpectedly found a row in a table with less then 2 cells")
                print(unique, l, path, i)
                continue
            
            if l == 2:
                case_study[unique[0]] = unique[1]
            else:
                if in_header:
                    in_header = False
                    case_study['column_names'] = unique
                else:
                    rows.append(unique)
            
        case_study["rows"] = rows
        case_studies.append(case_study)

    return case_studies

def parse_docx_to_dict_v2(path):
    document_docx = docx.Document(path)

    case_studies = []
    for i, table in enumerate(document_docx.tables):
        case_study = {
            "id": str(uuid.uuid4()),
            "case_id_in_doc": str(i),
            "docx_path": path
        }

        in_header = True
        provision_plan = []
        provision_keys = ["provision", "frequency", "delivered_by"]

        for row in table.rows:
            contents = [cell.text.strip() for cell in row.cells]
            unique = []
            [unique.append(x) for x in contents if x not in unique]
            l = len(unique)

            if l < 2:
                print("Unexpectedly found a row in a table with less than 2 cells")
                print(unique, l, path, i)
                continue

            # Case: Two cells → key-value pair (e.g. Special Educational Needs, Suggested Outcomes)
            if l == 2 and in_header:
                case_study[unique[0]] = unique[1]

            # Detect the start of provision plan by matching header row
            elif all(key.lower() in [c.lower() for c in unique] for key in ["Provision", "How frequently will it be delivered?"]):
                in_header = False  # now we are in the provision plan section

            # Actual provision plan rows (must have at least 3 values)
            elif not in_header and len(unique) >= 3:
                provision_entry = {
                    provision_keys[0]: unique[0],
                    provision_keys[1]: unique[1],
                    provision_keys[2]: unique[2]
                }
                provision_plan.append(provision_entry)

        case_study["provision_plan"] = provision_plan
        case_studies.append(case_study)

    return case_studies

                
def format_to_markdown(case_study):

    header = "| " + " | ".join(case_study["column_names"]) + " |"
    separator = "|" + "|".join(["-" * (len(col) + 2) for col in case_study["column_names"]]) + "|"

    markdown_table = header + "\n" + separator + "\n"

    for row in case_study["rows"]:
        cleaned_cells = [cell.replace("\n", " ").replace("\r", " ").replace("|", "\\|") for cell in row]
        row_str = "| " + " | ".join(cleaned_cells) + " |"
        markdown_table += row_str + "\n"

    print(markdown_table)

    final_text = f"""Anonymus personalised educational plan
    Special Educational Needs:
    {case_study['Special Educational Needs']}


    Suggested Outcome(s) and when you expect this to be achieved:
    {case_study['Suggested Outcome(s) and when you expect this to be achieved']}


    {markdown_table}    
    """
    return final_text


def case_study_json_to_text(json_case_study):
    text = "Case Study:\n\n"
    text += "Special Educational Needs: " + json_case_study.get("Special Educational Needs", "N/A") + "\n\n"
    text += "Suggested Outcome(s): " + json_case_study.get("Suggested Outcome(s) and when you expect this to be achieved", "N/A") + "\n\n"
    text += "Provision Plan:\n"
    for provision in json_case_study.get("provision_plan", []):
        line = f"- {provision.get('provision', '')}, The frequency would be: {provision.get('frequency', '')}, delivered by {provision.get('delivered_by', '')}"
        text += line + "\n"
    return text

# Get a ChromaDB collection or recreate it if specified 
def get_or_create_collection(client, name, metadata=None, config=None, recreate=False):
    existing_collections = [c for c in client.list_collections()]

    if name in existing_collections:
        if recreate:
            print(f"Collection '{name}' exists. Deleting and recreating...")
            client.delete_collection(name)
            coll_cs = parse_and_write_case_studies(name, metadata, config)
            return coll_cs
             
        else:
            print(f"Collection '{name}' exists. Returning existing collection...")
            return client.get_collection(name)
    else:
        print(f"Collection '{name}' does not exist. Creating new collection...")
        coll_cs = parse_and_write_case_studies(name, metadata, config)
        return coll_cs


def parse_and_write_case_studies(name, metadata=None, config=None):
    # model_name = 'sentence-transformers/all-mpnet-base-v2'
    embedder = SentenceTransformersDocumentEmbedder(model_name)
    embedder.warm_up()

    # coll_sc = client.get_or_create_collection("SOP_case_studies") #,  configuration={"hnsw": {"space": "cosine"}})
    coll_cs = client.create_collection(name=name, metadata=metadata, configuration=config)
   
    case_studies_dir = os.path.join(data_folder, "case_studies")
    case_studies_files = os.listdir(case_studies_dir)
    case_studies_paths = [os.path.join(case_studies_dir, file) for file in case_studies_files]

    case_studies = []
    for path in case_studies_paths:
        case_studies_in_dox = parse_docx_to_dict_v2(path)
        case_studies.extend(case_studies_in_dox)

    full_text_case_study = [case_study_json_to_text(cs) for cs in case_studies]

    haystack_docs = [
    Document(content=text, meta={
        "file": cs["docx_path"],
        "table_in_file": cs["case_id_in_doc"],
        "docx_path": cs["docx_path"],
        "case_id": cs["id"]
    })
    for text, cs in zip(full_text_case_study, case_studies)
]

    # --- EMBED DOCUMENTS ---
    docs_with_embeddings = embedder.run(haystack_docs)["documents"]
    embeddings = [doc.embedding for doc in docs_with_embeddings]

    # --- ADD TO CHROMA ---
    ids = [cs["id"] for cs in case_studies]
    metas = [doc.meta for doc in docs_with_embeddings]

    coll_cs.add(
        ids=ids,
        embeddings=embeddings,
        metadatas=metas,
        documents=full_text_case_study
)    
    return coll_cs

# %% 
# if __name__ == "__main__":
#     if pdf_kb:
#         document_store_pdf = create_pdf_kb()
#         document_store_pdf.save_to_disk("data/doc_store_pdfs.pkl")

#     if sc_kb:
#         document_store_scenario = create_scenario_kb()
#         document_store_scenario.save_to_disk("data/doc_store_scenarios.pkl")
        
#     if comb_kb:
#         document_store_combined = create_combined_kb() 
#         document_store_pdf.save_to_disk("data/doc_store_combined.pkl")

#     if case_studies_kb:
#         # parse_and_write_case_studies()

#         # config_obj = CollectionConfig(hnsw={"space": "cosine", "ef_construction": 250, "M": 18})

#         coll_cs = get_or_create_collection(client,
#                                  name="SOP_7_case_studies", 
#                                  metadata={'embedding_model':model_name},
#                                  recreate=False)
        

      
# %%
# document_store_pdf = preprocess_pdfs(InMemoryDocumentStore(embedding_similarity_function="cosine"))

# name = 'compliance'

# metadata = None
# config = None
# #%%
# name = 'compliance'
# coll_com = client.create_collection(name=name)
#%%
print("start printing")
print(doc_class)
print(docs_folder)
print(meta_docs_urls)
print("end of printing")
document_store = ChromaDocumentStore("compliance", host="localhost", port=8077)

# document_store_pdf = preprocess_pdfs(document_store)



#%%
print(doc_class)


#%%

